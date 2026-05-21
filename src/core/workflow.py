from collections import OrderedDict
from collections.abc import Callable

import docx

from src.core.excel_writer import ExcelWorkbookWriter
from src.core.models import MaterialGroup, MaterialItem
from src.core.profiles import MaterialProfileParser

ProgressCallback = Callable[[int, int, str], None]


class MaterialAutomationService:
    def __init__(self) -> None:
        self.parser = MaterialProfileParser()

    def _report(self, callback: ProgressCallback | None, current: int, total: int, message: str) -> None:
        if callback:
            callback(current, total, message)

    def extrair_dados_word(self, caminho_arquivo_word: str, callback: ProgressCallback | None = None, start_step: int = 0, total_steps: int = 1) -> list[MaterialItem] | None:
        documento = docx.Document(caminho_arquivo_word)
        try:
            tabela = documento.tables[0]
        except IndexError:
            return None

        if len(tabela.rows) < 2:
            return None

        perfis_str = tabela.cell(1, 0).text
        acos_str = tabela.cell(1, 1).text
        ltotais_str = tabela.cell(1, 2).text
        pesos_str = tabela.cell(1, 3).text

        lista_perfis = list(filter(None, perfis_str.strip().split("\n")))
        lista_acos = list(filter(None, acos_str.strip().split("\n")))
        lista_ltotais = list(filter(None, ltotais_str.strip().split("\n")))
        lista_pesos = list(filter(None, pesos_str.strip().split("\n")))

        num_perfis = len(lista_perfis)
        if num_perfis == 0:
            return None

        dados_finais: list[MaterialItem] = []

        for indice, perfil in enumerate(lista_perfis):
            perfil_limpo = perfil.strip()
            aco = lista_acos[indice].strip() if indice < len(lista_acos) else (lista_acos[0].strip() if lista_acos else "A36")

            comprimento_m = 0.0
            valor_coluna = lista_ltotais[indice].strip() if indice < len(lista_ltotais) else ""
            if valor_coluna:
                try:
                    comprimento_m = float(valor_coluna.replace(",", ".")) / 100.0
                except ValueError:
                    comprimento_m = 0.0

            if comprimento_m == 0.0:
                comprimento_m = self.parser.extrair_comprimento_texto(perfil_limpo)

            if "CA " in perfil_limpo.upper():
                comprimento_m *= 2

            peso_str = lista_pesos[indice].strip().replace(",", ".") if indice < len(lista_pesos) else "0"
            try:
                peso_final = float(peso_str)
            except ValueError:
                peso_final = 0.0

            dados_finais.append(self.parser.criar_item(perfil_limpo, aco, comprimento_m, peso_final))
            self._report(callback, start_step + indice + 1, total_steps, f"Extraindo dados ({indice + 1}/{num_perfis})")

        return dados_finais

    def agrupar_materiais(self, materiais: list[MaterialItem]) -> list[MaterialGroup]:
        grupos: "OrderedDict[str, MaterialGroup]" = OrderedDict()

        for item in materiais:
            grupo = grupos.get(item.chave_agrupamento)
            if grupo is None:
                grupos[item.chave_agrupamento] = MaterialGroup.from_item(item)
            else:
                grupo.merge(item)

        return list(grupos.values())

    def preencher_planilha_excel(self, caminho_planilha: str, dados_materiais: list[MaterialItem], callback: ProgressCallback | None = None, start_step: int = 0, total_steps: int = 1) -> str:
        writer = ExcelWorkbookWriter(caminho_planilha, parser=self.parser)
        grupos = self.agrupar_materiais(dados_materiais)
        linha_de_busca_por_secao: dict[str, int] = {}

        self._report(callback, start_step, total_steps, "Agrupando materiais...")

        for indice, grupo in enumerate(grupos):
            if grupo.tipo_perfil == "VIGA_W":
                linha_alvo = writer.encontrar_linha_viga_w(grupo.nome_normalizado)
            else:
                linha_inicio = linha_de_busca_por_secao.get(grupo.codigo_excel, 4)
                linha_alvo = writer.encontrar_proxima_linha_vazia(grupo.codigo_excel, linha_inicio)

            if linha_alvo is not None:
                writer.escrever_grupo(linha_alvo, grupo)
                if grupo.tipo_perfil != "VIGA_W":
                    linha_de_busca_por_secao[grupo.codigo_excel] = linha_alvo + 1

            self._report(callback, start_step + indice + 1, total_steps, f"Preenchendo planilha ({indice + 1}/{len(grupos)})")

        writer.ocultar_linhas_vazias()
        self._report(callback, start_step + len(grupos) + 1, total_steps, "Salvando arquivo processado...")
        return writer.salvar()

    def processar(self, caminho_arquivo_word: str, caminho_planilha: str, callback: ProgressCallback | None = None) -> str | None:
        documento = docx.Document(caminho_arquivo_word)
        try:
            tabela = documento.tables[0]
        except IndexError:
            return None

        if len(tabela.rows) < 2:
            return None

        quantidade_itens = len(list(filter(None, tabela.cell(1, 0).text.strip().split("\n"))))
        total_steps = max(quantidade_itens * 2 + 2, 1)

        self._report(callback, 0, total_steps, "Abrindo documento Word...")
        dados_materiais = self.extrair_dados_word(caminho_arquivo_word, callback=callback, start_step=0, total_steps=total_steps)
        if not dados_materiais:
            return None

        caminho_processado = self.preencher_planilha_excel(caminho_planilha, dados_materiais, callback=callback, start_step=quantidade_itens + 1, total_steps=total_steps)
        self._report(callback, total_steps, total_steps, "Processamento concluido.")
        return caminho_processado
